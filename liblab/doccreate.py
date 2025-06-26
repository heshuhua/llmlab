from docx import Document
import os

def create_doc(filename, content, doc_type, date):
    doc = Document()
    doc.add_heading(filename.replace(".docx", ""), level=1)
    doc.add_paragraph(content)
    # 可以在文档属性中添加更多元数据，但这里为了简单，我们主要通过文件名和外部元数据管理
    # 也可以在文档内容中包含一些结构化信息，比如：
    doc.add_paragraph(f"\n文档类型: {doc_type}")
    doc.add_paragraph(f"创建日期: {date}")
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    if not os.path.exists("word_kb"):
        os.makedirs("word_kb")

    # 文档1：劳动合同试用期条款
    content1 = (
        "劳动合同试用期规定：根据《中华人民共和国劳动合同法》第十九条，"
        "劳动合同期限三个月以上不满一年的，试用期不得超过一个月；"
        "劳动合同期限一年以上不满三年的，试用期不得超过二个月；"
        "三年以上固定期限和无固定期限的劳动合同，试用期不得超过六个月。"
        "同一用人单位与同一劳动者只能约定一次试用期。"
    )
    create_doc("word_kb/劳动合同-试用期规定.docx", content1, "合同模板", "2024-01-15")

    # 文档2：服务合同违约条款
    content2 = (
        "服务合同违约责任：如果一方违反本合同的任何条款，违约方应向守约方支付违约金，"
        "违约金金额为合同总价的10%。若违约金不足以弥补守约方损失的，违约方还应赔偿差额损失。"
        "因不可抗力导致无法履行合同的，不视为违约。"
    )
    create_doc("word_kb/服务合同-违约责任.docx", content2, "合同模板", "2023-11-20")

    # 文档3：知识产权法律意见书
    content3 = (
        "关于软件著作权保护的法律意见：软件著作权自开发完成之日起产生，"
        "受《中华人民共和国著作权法》保护。建议及时进行软件著作权登记，"
        "以确权并作为维权的初步证据。侵犯软件著作权的行为包括未经许可复制、发行、出租等。"
    )
    create_doc("word_kb/知识产权-软件著作权保护.docx", content3, "法律意见书", "2024-03-01")

    # 文档4：员工请假流程
    content4 = (
        "员工请假流程：员工需提前通过内部OA系统提交请假申请，并附上相关证明材料（如病假条）。"
        "请假审批权限：3天以内由部门经理审批，3天以上由总监审批，15天以上需报总经理审批。"
    )
    create_doc("word_kb/内部规章-请假流程.docx", content4, "内部规章", "2023-10-01")

    print("\n示例 Word 文档创建完成。")